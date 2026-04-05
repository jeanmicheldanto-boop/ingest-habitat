"""
Générateur de courriers de prospection personnalisés au format Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os


def add_border_bottom(paragraph):
    """Ajoute une bordure en bas d'un paragraphe"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_prospection_letter(recipient_data: dict, output_path: str, test_mode: bool = False):
    """
    Crée un courrier de prospection personnalisé
    
    Args:
        recipient_data: Dictionnaire avec les données du destinataire
        output_path: Chemin du fichier de sortie
        test_mode: Si True, affiche les données utilisées
    """
    
    doc = Document()
    
    # Marges réduites pour tenir sur 1 page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # === EN-TÊTE : Table avec 2 colonnes ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Colonne gauche : Expéditeur
    left_cell = header_table.rows[0].cells[0]
    left_cell.width = Inches(3.5)
    
    p = left_cell.paragraphs[0]
    p.add_run("Patrick Danto\n").bold = True
    p.add_run("Associé BMSE\n")
    p.add_run("Directeur technique ConfidensIA\n")
    p.add_run("55 rue de l'Abbé Carton\n")
    p.add_run("75014 Paris")
    
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    # Colonne droite : Logo + Destinataire
    right_cell = header_table.rows[0].cells[1]
    right_cell.width = Inches(3.5)
    
    # Logo BMSE
    logo_path = "__pycache__/logo_bmse.png"
    if os.path.exists(logo_path):
        p_logo = right_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    else:
        p_logo = right_cell.paragraphs[0]
        p_logo.add_run("[Logo BMSE]").bold = True
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Adresse destinataire
    p_dest = right_cell.add_paragraph()
    p_dest.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    civilite = recipient_data.get('civilite', 'Madame, Monsieur')
    nom = recipient_data.get('nom', '')
    titre = recipient_data.get('titre', '')
    organisme = recipient_data.get('organisme', '')
    adresse = recipient_data.get('adresse', '')
    
    if nom:
        p_dest.add_run(f"{civilite} {nom}\n").bold = True
    if titre:
        p_dest.add_run(f"{titre}\n")
    if organisme:
        p_dest.add_run(f"{organisme}\n").bold = True
    if adresse:
        p_dest.add_run(f"{adresse}")
    
    for run in p_dest.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    # === OBJET ===
    p_objet = doc.add_paragraph()
    p_objet.add_run("Objet : ").bold = True
    p_objet.add_run("Intelligence artificielle et données dans les ESSMS : un accompagnement ancré dans vos réalités").bold = True
    p_objet.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_objet.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # === CORPS DU TEXTE ===
    
    # Salutation
    p_salut = doc.add_paragraph()
    salutation = f"{civilite} {nom.split()[-1] if nom else ''},"
    p_salut.add_run(salutation)
    p_salut.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p_salut.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 1
    p1 = doc.add_paragraph()
    p1.add_run("Les établissements et services sociaux et médico-sociaux font face à une inflation des tâches administratives, des obligations de reporting et des exigences de conformité, dans un contexte de ressources déjà tendues.")
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p1.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 2
    p2 = doc.add_paragraph()
    p2.add_run("L'intelligence artificielle ne bouleversera pas les fondamentaux des métiers de l'humain et du travail social. Elle n'est pas une réponse à tous les problèmes. Mais ")
    p2.add_run("mobilisée avec méthode et responsabilité, elle peut libérer du temps administratif au profit de l'accompagnement.").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p2.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 3
    p3 = doc.add_paragraph()
    p3.add_run("BMSE, cabinet spécialisé dans la tarification et l'accompagnement budgétaire des ESSMS, a développé une offre spécifique en intelligence artificielle et analyse de données, en partenariat avec ConfidensIA, entreprise française de R&D qui conçoit des outils d'IA pour le secteur social et médico-social.")
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p3.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 4
    p4 = doc.add_paragraph()
    p4.add_run("Notre spécificité : ")
    p4.add_run("nous connaissons à la fois vos réalités de terrain et ce que peut apporter la transformation numérique dans votre quotidien.").bold = True
    p4.add_run(" Nous accompagnons déjà plusieurs organismes gestionnaires et autorités de tarification.")
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 5
    p5 = doc.add_paragraph()
    p5.add_run("Nos solutions personnalisées portent sur :")
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p5.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Liste à puces
    bullet1 = doc.add_paragraph(style='List Bullet')
    bullet1.add_run("L'automatisation des tâches de reporting et d'analyse financière").bold = True
    bullet1.add_run(" via le Hub des ESSMS (extraction et analyse de vos données budgétaires issues des cadres normalisés)")
    bullet1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet1.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    bullet2.add_run("L'anonymisation et la pseudonymisation des écrits professionnels").bold = True
    bullet2.add_run(" pour permettre un usage conforme au RGPD de l'IA générative dans vos pratiques rédactionnelles")
    bullet2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet2.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    bullet3.add_run("L'accompagnement stratégique").bold = True
    bullet3.add_run(" pour définir une doctrine d'usage de l'IA qui colle à vos valeurs, identifier les cas d'usage prioritaires et former vos équipes aux principes fondamentaux et aux risques de l'IA")
    bullet3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet3.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 6
    p6 = doc.add_paragraph()
    p6.add_run("Une bonne manière de commencer consiste souvent en une ")
    p6.add_run("formation-action courte (½ journée à 2 jours selon vos besoins)").bold = True
    p6.add_run(" permettant de diffuser les principes de l'IA, d'en mesurer les limites, de poser collectivement une charte d'usage et d'envisager des cas d'usage concrets et utiles.")
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p6.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 7
    p7 = doc.add_paragraph()
    p7.add_run("Vous trouverez en pièce jointe une présentation de notre offre.")
    p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p7.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 8
    p8 = doc.add_paragraph()
    p8.add_run("Je reste à votre disposition pour en échanger sans engagement de votre part.")
    p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p8.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Formule de politesse
    p_pol = doc.add_paragraph()
    p_pol.add_run("Bien cordialement,")
    p_pol.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_pol.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    doc.add_paragraph()

    # Texte de signature (texte conservé; image supprimée pour signature manuscrite)
    p_sig = doc.add_paragraph()
    p_sig.add_run("Patrick Danto\n").bold = True
    p_sig.add_run("patrick.danto@bmse.fr\n")
    p_sig.add_run("BMSE et ConfidensIA")
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_sig.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Sauvegarder
    doc.save(output_path)
    
    if test_mode:
        print(f"✅ Courrier créé : {output_path}")
        print(f"   Destinataire : {civilite} {nom}")
        print(f"   Organisme : {organisme}")


def generate_test_letter():
    """Génère un courrier de test avec le premier destinataire du fichier"""
    
    # Charger le fichier de publipostage
    df = pd.read_excel('outputs/prospection_250_FINAL_FORMATE_V2_PUBLIPOSTAGE.xlsx')
    
    # Premier destinataire
    first = df.iloc[0]
    
    recipient_data = {
        'civilite': first.get('Civilité', 'Madame, Monsieur'),
        'nom': first.get('dirigeant_nom', ''),
        'titre': first.get('dirigeant_titre', ''),
        'organisme': first.get('nom_public', ''),
        'adresse': first.get('Adresse Publipostage', '')
    }
    
    output_path = 'outputs/courrier_test_v4.docx'
    
    print("=" * 70)
    print("📄 GÉNÉRATION COURRIER DE TEST")
    print("=" * 70)
    print(f"\nDonnées utilisées :")
    for key, value in recipient_data.items():
        print(f"  {key:12s}: {value}")
    print()
    
    create_prospection_letter(recipient_data, output_path, test_mode=True)
    
    print(f"\n✅ Courrier de test créé : {output_path}")
    print("\nVous pouvez l'ouvrir avec Word pour vérifier le rendu.")


def generate_all_letters():
    """Génère tous les courriers dans un seul fichier Word"""
    
    # Charger le fichier de publipostage
    df = pd.read_excel('outputs/prospection_250_FINAL_FORMATE_V2_PUBLIPOSTAGE.xlsx')
    
    output_path = 'outputs/courriers_prospection_250.docx'
    
    print("=" * 70)
    print("📄 GÉNÉRATION DE TOUS LES COURRIERS")
    print("=" * 70)
    print(f"\nNombre de destinataires : {len(df)}")
    print(f"Fichier de sortie : {output_path}\n")
    
    # Créer le document
    doc = Document()
    
    # Configurer les marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Générer chaque courrier
    for idx, row in df.iterrows():
        if idx > 0:
            # Saut de page avant chaque nouveau courrier (sauf le premier)
            doc.add_page_break()
        
        recipient_data = {
            'civilite': row.get('Civilité', 'Madame, Monsieur'),
            'nom': row.get('dirigeant_nom', ''),
            'titre': row.get('dirigeant_titre', ''),
            'organisme': row.get('nom_public', ''),
            'adresse': row.get('Adresse Publipostage', '')
        }
        
        # Ajouter le courrier au document
        add_letter_to_document(doc, recipient_data)
        
        # Afficher la progression tous les 10 courriers
        if (idx + 1) % 10 == 0:
            print(f"  ✓ {idx + 1}/{len(df)} courriers générés...")
    
    # Sauvegarder le document
    doc.save(output_path)
    
    print(f"\n✅ Tous les courriers ont été générés : {output_path}")
    print(f"   Total : {len(df)} courriers")


def add_letter_to_document(doc, recipient_data):
    """Ajoute un courrier au document existant"""
    
    # === EN-TÊTE : Table avec 2 colonnes ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Colonne gauche : Expéditeur
    left_cell = header_table.rows[0].cells[0]
    left_cell.width = Inches(3.5)
    
    p = left_cell.paragraphs[0]
    p.add_run("Patrick Danto\n").bold = True
    p.add_run("Associé BMSE\n")
    p.add_run("Directeur technique ConfidensIA\n")
    p.add_run("55 rue de l'Abbé Carton\n")
    p.add_run("75014 Paris")
    
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    # Colonne droite : Logo + Destinataire
    right_cell = header_table.rows[0].cells[1]
    right_cell.width = Inches(3.5)
    
    # Logo BMSE
    logo_path = "__pycache__/logo_bmse.png"
    if os.path.exists(logo_path):
        p_logo = right_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    else:
        p_logo = right_cell.paragraphs[0]
        p_logo.add_run("[Logo BMSE]").bold = True
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Adresse destinataire
    p_dest = right_cell.add_paragraph()
    p_dest.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    civilite = recipient_data.get('civilite', 'Madame, Monsieur')
    nom = recipient_data.get('nom', '')
    titre = recipient_data.get('titre', '')
    organisme = recipient_data.get('organisme', '')
    adresse = recipient_data.get('adresse', '')
    
    if nom:
        p_dest.add_run(f"{civilite} {nom}\n").bold = True
    if titre:
        p_dest.add_run(f"{titre}\n")
    if organisme:
        p_dest.add_run(f"{organisme}\n").bold = True
    if adresse:
        p_dest.add_run(f"{adresse}")
    
    for run in p_dest.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    # === OBJET ===
    p_objet = doc.add_paragraph()
    p_objet.add_run("Objet : ").bold = True
    p_objet.add_run("Intelligence artificielle et données dans les ESSMS : un accompagnement ancré dans vos réalités").bold = True
    p_objet.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_objet.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # === CORPS DU TEXTE ===
    
    # Salutation
    p_salut = doc.add_paragraph()
    salutation = f"{civilite} {nom.split()[-1] if nom else ''},"
    p_salut.add_run(salutation)
    p_salut.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p_salut.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 1
    p1 = doc.add_paragraph()
    p1.add_run("Les établissements et services sociaux et médico-sociaux font face à une inflation des tâches administratives, des obligations de reporting et des exigences de conformité, dans un contexte de ressources déjà tendues.")
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p1.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 2
    p2 = doc.add_paragraph()
    p2.add_run("L'intelligence artificielle ne bouleversera pas les fondamentaux des métiers de l'humain et du travail social. Elle n'est pas une réponse à tous les problèmes. Mais ")
    p2.add_run("mobilisée avec méthode et responsabilité, elle peut libérer du temps administratif au profit de l'accompagnement.").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p2.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 3
    p3 = doc.add_paragraph()
    p3.add_run("BMSE, cabinet spécialisé dans la tarification et l'accompagnement budgétaire des ESSMS, a développé une offre spécifique en intelligence artificielle et analyse de données, en partenariat avec ConfidensIA, entreprise française de R&D qui conçoit des outils d'IA pour le secteur social et médico-social.")
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p3.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 4
    p4 = doc.add_paragraph()
    p4.add_run("Notre spécificité : ")
    p4.add_run("nous connaissons à la fois vos réalités de terrain et ce que peut apporter la transformation numérique dans votre quotidien.").bold = True
    p4.add_run(" Nous accompagnons déjà plusieurs organismes gestionnaires et autorités de tarification.")
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 5
    p5 = doc.add_paragraph()
    p5.add_run("Nos solutions personnalisées portent sur :")
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p5.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Liste à puces
    bullet1 = doc.add_paragraph(style='List Bullet')
    bullet1.add_run("L'automatisation des tâches de reporting et d'analyse financière").bold = True
    bullet1.add_run(" via le Hub des ESSMS (extraction et analyse de vos données budgétaires issues des cadres normalisés)")
    bullet1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet1.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    bullet2.add_run("L'anonymisation et la pseudonymisation des écrits professionnels").bold = True
    bullet2.add_run(" pour permettre un usage conforme au RGPD de l'IA générative dans vos pratiques rédactionnelles")
    bullet2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet2.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    bullet3.add_run("L'accompagnement stratégique").bold = True
    bullet3.add_run(" pour définir une doctrine d'usage de l'IA qui colle à vos valeurs, identifier les cas d'usage prioritaires et former vos équipes aux principes fondamentaux et aux risques de l'IA")
    bullet3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in bullet3.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 6
    p6 = doc.add_paragraph()
    p6.add_run("Une bonne manière de commencer consiste souvent en une ")
    p6.add_run("formation-action courte (½ journée à 2 jours selon vos besoins)").bold = True
    p6.add_run(" permettant de diffuser les principes de l'IA, d'en mesurer les limites, de poser collectivement une charte d'usage et d'envisager des cas d'usage concrets et utiles.")
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p6.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 7
    p7 = doc.add_paragraph()
    p7.add_run("Vous trouverez en pièce jointe une présentation de notre offre.")
    p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p7.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    # Paragraphe 8
    p8 = doc.add_paragraph()
    p8.add_run("Je reste à votre disposition pour en échanger sans engagement de votre part.")
    p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p8.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Formule de politesse
    p_pol = doc.add_paragraph()
    p_pol.add_run("Bien cordialement,")
    p_pol.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_pol.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Texte de signature (texte conservé; image supprimée pour signature manuscrite)
    p_sig = doc.add_paragraph()
    p_sig.add_run("Patrick Danto\n").bold = True
    p_sig.add_run("patrick.danto@bmse.fr\n")
    p_sig.add_run("BMSE et ConfidensIA")
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_sig.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "--all":
        generate_all_letters()
    else:
        generate_test_letter()
