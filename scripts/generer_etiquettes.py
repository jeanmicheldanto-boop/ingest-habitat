"""
Générateur d'étiquettes d'adresse au format Avery L7163
14 étiquettes par page (2 colonnes x 7 lignes)
Format : 99.1 x 38.1 mm
"""

from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd


def create_label_table(doc):
    """Crée un tableau pour une page d'étiquettes Avery L7163"""
    
    # 14 étiquettes = 7 lignes x 2 colonnes
    table = doc.add_table(rows=7, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Dimensions Avery L7163
    label_width = Mm(99.1)
    label_height = Mm(37.0)  # Légèrement réduit pour tenir sur la page
    
    # Configurer chaque cellule
    for row in table.rows:
        row.height = label_height
        for idx, cell in enumerate(row.cells):
            cell.width = label_width
            # Marges intérieures
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            # Marges intérieures : augmenter légèrement la marge gauche
            # top/bottom/right = 1.5mm (≈85 twips), left = 2.0mm (≈113 twips)
            for margin_name in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin_name}')
                if margin_name == 'left':
                    node.set(qn('w:w'), '170')  # ~3.0mm en twips (augmenté pour décaler à droite)
                else:
                    node.set(qn('w:w'), '85')   # ~1.5mm en twips
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            
            tcPr.append(tcMar)
    
    # Définir la largeur exacte des colonnes
    table.columns[0].width = label_width
    table.columns[1].width = label_width
    
    # Supprimer les bordures
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # Supprimer l'espacement entre les cellules
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Espacement cellules = 0
    tblCellSpacing = OxmlElement('w:tblCellSpacing')
    tblCellSpacing.set(qn('w:w'), '0')
    tblCellSpacing.set(qn('w:type'), 'dxa')
    tblPr.append(tblCellSpacing)
    
    # Espacement entre lignes = 0 (très important !)
    for row in table.rows:
        trPr = row._element.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(label_height.twips)))
        trHeight.set(qn('w:hRule'), 'exact')  # Hauteur EXACTE
        trPr.append(trHeight)
    
    return table


def add_address_to_cell(cell, civilite, nom, titre, organisme, adresse):
    """Ajoute une adresse dans une cellule d'étiquette"""
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Décalage interne à gauche pour toutes les lignes d'adresse
    p_format = p.paragraph_format
    p_format.left_indent = Mm(3.0)
    
    # Supprimer TOUS les espacements de paragraphe
    pPr = p._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')  # Interligne simple (12pt = 240 twips)
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    
    # Civilité + Nom
    if nom:
        run = p.add_run(f"{civilite} {nom}\n")
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.bold = True
    
    # Titre
    if titre and str(titre) != 'nan':
        run = p.add_run(f"{titre}\n")
        run.font.size = Pt(9)
        run.font.name = 'Arial'
    
    # Organisme
    if organisme:
        run = p.add_run(f"{organisme}\n")
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.bold = True
    
    # Adresse
    if adresse:
        run = p.add_run(f"{adresse}")
        run.font.size = Pt(9)
        run.font.name = 'Arial'


def generate_labels():
    """Génère toutes les étiquettes d'adresse"""
    
    # Charger le fichier de publipostage
    df = pd.read_excel('outputs/prospection_250_FINAL_FORMATE_V2_PUBLIPOSTAGE.xlsx')
    
    output_path = 'outputs/etiquettes_prospection_250_v7.docx'
    
    print("=" * 70)
    print("🏷️  GÉNÉRATION DES ÉTIQUETTES D'ADRESSE")
    print("=" * 70)
    print(f"\nFormat : Avery L7163 (14 étiquettes/page)")
    print(f"Nombre de destinataires : {len(df)}")
    print(f"Nombre de pages : {(len(df) + 13) // 14}\n")
    
    # Créer le document
    doc = Document()
    
    # Configurer les marges (Avery L7163 specs officielles)
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(13.0)   # Marge haute réduite
        section.bottom_margin = Mm(13.0) # Marge basse réduite
        section.left_margin = Mm(7.0)   # Marge gauche légèrement augmentée
        section.right_margin = Mm(5.5)  # Marge droite
        section.page_height = Mm(297)   # A4
        section.page_width = Mm(210)    # A4
    
    labels_on_page = 0
    current_table = None
    current_row = 0
    current_col = 0
    
    for idx, row in df.iterrows():
        # Nouvelle page tous les 14 labels
        if labels_on_page == 0:
            # Pas de page_break - on laisse les tableaux se suivre naturellement
            current_table = create_label_table(doc)
            current_row = 0
            current_col = 0
        
        # Données du destinataire
        civilite = row.get('Civilité', 'Madame, Monsieur')
        nom = row.get('dirigeant_nom', '')
        titre = row.get('dirigeant_titre', '')
        organisme = row.get('nom_public', '')
        adresse = row.get('Adresse Publipostage', '')
        
        # Ajouter l'adresse dans la cellule
        cell = current_table.rows[current_row].cells[current_col]
        add_address_to_cell(cell, civilite, nom, titre, organisme, adresse)
        
        # Progression
        labels_on_page += 1
        current_col += 1
        if current_col == 2:
            current_col = 0
            current_row += 1
        
        # Reset après 14 étiquettes
        if labels_on_page == 14:
            labels_on_page = 0
        
        # Afficher progression
        if (idx + 1) % 28 == 0:
            print(f"  ✓ {idx + 1}/{len(df)} étiquettes générées...")
    
    # Sauvegarder
    doc.save(output_path)
    
    print(f"\n✅ Étiquettes générées : {output_path}")
    print(f"   Total : {len(df)} étiquettes sur {(len(df) + 13) // 14} pages")
    print(f"\n💡 Impression recommandée : Avery L7163 (14 étiquettes/page)")


if __name__ == "__main__":
    generate_labels()
